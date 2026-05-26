# 后端总入口。
# 这个文件把后端几条主链路都串在一起：
# 1. 文档解析与上传；
# 2. 图片 / 视频证据分析；
# 3. 会话级检索与知识图谱接口；
# 4. 传感器数据接入；
# 5. 普通聊天与多智能体问答。
from flask import Flask, request, jsonify
from flask_cors import CORS
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import re
import os
import tempfile
import requests
import base64
import mimetypes
from pathlib import Path
import json
from io import BytesIO
from config import config
from retrieval import (
    hydrate_session_documents,
    ingest_document,
    retrieve_relevant_chunks,
    has_session_documents,
    remove_document,
    list_session_documents,
    list_session_chunks,
)
from sensor_store import push_session_sensors, list_session_sensors, has_session_sensors, clear_session_sensors, remove_session_sensor
from knowledge_graph import (
    build_and_store_session_graph,
    get_session_graph,
    query_graph,
    query_centered_graph,
    clear_session_graph,
    expand_graph_neighbors,
    start_graph_build,
    get_graph_build_status,
    mark_graph_build_pending,
    import_triples_graph,
)
from risk_fusion import build_risk_profile
from persistence import (
    init_storage,
    register_user,
    login_user,
    get_user_by_token,
    delete_token,
    save_document_asset,
    list_document_assets,
    delete_document_asset,
    save_image_asset,
    list_image_assets,
    delete_image_asset,
    save_video_asset,
    list_video_assets,
    delete_video_asset,
    save_sensor_records,
    list_sensor_records,
    clear_sensor_records,
    delete_sensor_record,
    save_message,
    list_messages,
)

# LangChain Agent
from agent import multi_agent_ask

# 创建整个后端共用的 Flask 应用实例。
app = Flask(__name__)
init_storage()

# 更明确的CORS配置
CORS(
    app,
    origins=config.CORS_ORIGINS,
    allow_headers=["Content-Type", "Authorization"],
    methods=["GET", "POST", "OPTIONS"],
)


# 从请求头中解析 Bearer Token。
def _extract_auth_token() -> str:
    # 从请求头 Authorization: Bearer <token> 中提取 token 字符串，失败则返回空串。
    header = str(request.headers.get("Authorization", "") or "").strip()
    if header.lower().startswith("bearer "):
        return header[7:].strip()
    return ""


# 校验当前请求的登录用户是否有效。
def _require_auth_user():
    # 验证当前请求的 Bearer token 有效，返回 (user_dict, token)。
    # token 无效则抛出 PermissionError，由上层接口返回 401。
    token = _extract_auth_token()
    user = get_user_by_token(token)
    if not user:
        raise PermissionError("未登录或登录已失效。")
    return user, token


# 解析并校验当前用户对应的知识库空间会话号。
def _resolve_user_session(requested_session_id: str | None = None):
    # 鉴权 + 校验前端传的 session_id 是否与登录用户的 library_session_id 一致。
    # 返回 (user_dict, stable_session_id, token)。
    user, token = _require_auth_user()
    stable_session_id = str(user.get("library_session_id") or "").strip()
    incoming_session_id = str(requested_session_id or "").strip()
    if incoming_session_id and incoming_session_id != stable_session_id:
        raise PermissionError("会话标识与当前登录用户不匹配。")
    return user, stable_session_id, token


# 把持久化数据重新灌回运行时缓存。
def _hydrate_persistent_runtime_state(user) -> None:
    # 登录后把持久化的文档和传感器数据恢复到进程内存中，使 RAG 检索和传感器缓存立即可用。
    user_id = int(user["id"])
    session_id = str(user["library_session_id"])

    persisted_docs = list_document_assets(user_id)
    if persisted_docs and not has_session_documents(session_id):
        hydrate_session_documents(session_id, persisted_docs)

    persisted_sensors = list_sensor_records(user_id)
    if persisted_sensors and not has_session_sensors(session_id):
        push_session_sensors(session_id=session_id, payload=persisted_sensors)


# 根据文件名猜测上传文件的 MIME 类型。
def _guess_upload_mime(file_name: str, default: str = "application/octet-stream") -> str:
    # 根据文件名后缀猜测 MIME 类型，用于持久化时记录文件类型。
    mime_type = mimetypes.guess_type(str(file_name or ""))[0]
    return mime_type or default


# 从原始字节流中提取文档文本内容。
def _extract_text_from_upload_bytes(file_name: str, raw_bytes: bytes):
    # 根据文件后缀从 bytes 中提取全文，支持 PDF/DOCX/TXT。
    # 返回 {"text": ..., "char_count": ...}（PDF 额外包含 "page_count"）。
    filename = str(file_name or "").lower()
    if filename.endswith(".pdf"):
        doc = fitz.open(stream=raw_bytes, filetype="pdf")
        full_text = ""
        for page_num in range(len(doc)):
            page = doc[page_num]
            full_text += page.get_text() + "\n"
        cleaned_text = clean_text(full_text)
        return {
            "text": cleaned_text,
            "char_count": len(cleaned_text),
            "page_count": len(doc),
        }
    if filename.endswith(".docx"):
        doc = DocxDocument(BytesIO(raw_bytes))
        text = "\n".join([para.text for para in doc.paragraphs])
        return {
            "text": text,
            "char_count": len(text),
        }
    if filename.endswith(".txt"):
        text = raw_bytes.decode("utf-8", errors="ignore")
        return {
            "text": text,
            "char_count": len(text),
        }
    raise ValueError("不支持的文件格式，请上传 TXT、DOCX 或 PDF")


# 统一整理图谱统计信息。
def _graph_payload(graph):
    # 统一整理图谱统计信息，避免每个接口都自己手工拼 node/relation 数量。
    stats = graph.get("stats") if isinstance(graph, dict) else {}
    return {
        "node_count": int((stats or {}).get("node_count") or len(graph.get("nodes", []))),
        "relation_count": int((stats or {}).get("relation_count") or len(graph.get("relations", []))),
        "stats": stats or {},
    }


# 用当前会话的检索片段重新构建知识图谱。
def _rebuild_session_knowledge_graph(session_id):
    # 用当前会话已经入库的检索片段重新构建知识图谱。
    chunks = list_session_chunks(session_id)
    graph = build_and_store_session_graph(session_id, chunks)
    return graph

VISION_ALLOWED_RISK_LEVELS = {"低", "中", "高", "极高", "未识别"}
VISION_PROMPT_TEMPLATE = (
    "你是煤矿安全应急图像识别助手。"
    "请严格围绕煤矿场景识别画面中的风险线索、作业场景和关键对象，"
    "不要回答与图像无关的内容。"
    "请只返回 JSON，对象格式为："
    '{"keywords":["关键词1","关键词2","关键词3"],'
    '"summary":"一句不超过30字的中文摘要",'
    '"risk_level":"低/中/高/极高/未识别"}。'
    "若无法判断，请返回 keywords 为空数组，summary 为“未识别”，risk_level 为“未识别”。"
)

# 清洗从 PDF 抽取出的原始文本。
def clean_text(text):
    """清洗PDF提取的文本"""
    # 1. 删除页眉页脚
    text = text.replace('应急管理部规章', '').replace('应急管理部发布', '')

    # 2. 删除页码 "- 数字 -"
    text = re.sub(r'-\s*\d+\s*-', '', text)

    # 3. 修复断句问题 - 非句末的换行改为空格
    text = re.sub(r'(?<![。；？！])\n', ' ', text)

    # 4. 规范化空格
    text = re.sub(r'\s+', ' ', text).strip()

    # 5. 确保"第"和条号之间没有多余空格
    text = re.sub(r'第\s+([一二三四五六七八九十百千万\d]+)\s*条', r'第\1条', text)

    return text


# 截断过长文本以控制上下文长度。
def _clip_text(text, max_len):
    # 对长文本做截断，主要用于：
    # 1. 控制提示词长度；
    # 2. 控制接口返回中的摘要长度；
    # 3. 避免日志或前端面板被超长文本撑爆。
    text = str(text or "").strip()
    if len(text) <= max_len:
        return text
    head = max_len // 2
    tail = max_len - head
    return f"{text[:head]}\n...(中间已省略)...\n{text[-tail:]}"


# 清洗并裁剪前端传来的历史对话。
def _normalize_history(history):
    # 把前端传来的历史对话清洗成后端统一格式，
    # 只保留 user/assistant 两类消息，并限制最大轮数。
    if not isinstance(history, list):
        return []

    normalized = []
    max_messages = max(2, config.AGENT_MAX_HISTORY_TURNS * 2)
    for item in history:
        if not isinstance(item, dict):
            continue
        role = str(item.get("role", "")).strip().lower()
        content = str(item.get("content", "")).strip()
        if role not in {"user", "assistant"} or not content:
            continue
        normalized.append({
            "role": role,
            "content": _clip_text(content, 600),
        })

    return normalized[-max_messages:]


# 清洗前端上传的文档证据结构。
def _normalize_document_evidence(documents):
    # 清洗文档证据，统一字段名、长度和来源类型，
    # 让后面的多智能体流程不必关心前端字段是否有别名。
    if not isinstance(documents, list):
        return []

    normalized = []
    for idx, item in enumerate(documents[:config.AGENT_MAX_EVIDENCE_DOCUMENTS]):
        if not isinstance(item, dict):
            continue
        text = str(item.get("text", "")).strip()
        if not text:
            continue
        doc_name = str(item.get("doc_name") or item.get("docName") or "未命名文档").strip() or "未命名文档"
        chunk_id = str(item.get("chunk_id") or item.get("chunkId") or f"{doc_name}:{idx}").strip()
        score = item.get("score")
        try:
            score = float(score) if score is not None else None
        except (TypeError, ValueError):
            score = None
        normalized.append({
            "doc_name": doc_name,
            "chunk_id": chunk_id,
            "text": _clip_text(text, 900),
            "score": score,
            "source_type": str(item.get("source_type") or "uploaded_doc"),
        })

    return normalized


# 清洗前端上传的图片证据结构。
def _normalize_image_evidence(images):
    # 图片、视频命中帧在后端最终都按“图像证据”来处理，
    # 这里负责把它们规整成统一结构。
    if not isinstance(images, list):
        return []

    normalized = []
    for item in images:
        if not isinstance(item, dict):
            continue
        summary = str(item.get("summary", "")).strip()
        if not summary:
            continue
        normalized.append({
            "image_name": str(item.get("image_name") or item.get("imageName") or "未命名图片").strip() or "未命名图片",
            "summary": _clip_text(summary, 300),
            "source_type": str(item.get("source_type") or "image_analysis"),
        })

    return normalized


# 清洗前端上传的传感器证据结构。
def _normalize_sensor_evidence(sensors):
    # 传感器记录同样先做一遍规整，避免不同来源的字段不一致。
    if not isinstance(sensors, list):
        return []

    normalized = []
    for idx, item in enumerate(sensors[:20]):
        if not isinstance(item, dict):
            continue
        sensor_id = str(item.get("sensor_id") or item.get("sensorId") or item.get("id") or f"sensor-{idx + 1}").strip()
        name = str(item.get("name") or item.get("sensor_name") or item.get("sensorName") or sensor_id).strip() or sensor_id
        value = item.get("value")
        unit = str(item.get("unit") or "").strip()
        threshold = item.get("threshold")
        timestamp = str(item.get("timestamp") or item.get("time") or "").strip()
        location = str(item.get("location") or item.get("area") or item.get("place") or "").strip()
        status = str(item.get("status") or item.get("state") or "").strip()
        source_type = str(item.get("source_type") or "sensor").strip() or "sensor"

        normalized.append({
            "sensor_id": sensor_id,
            "name": name,
            "value": value,
            "value_text": str(value).strip() if value is not None else "",
            "unit": unit,
            "threshold": threshold,
            "timestamp": timestamp,
            "location": location,
            "status": status,
            "source_type": source_type,
        })
    return normalized


# 把 agent-chat 请求统一整理成标准输入。
def _normalize_agent_chat_request(data):
    # agent-chat 是当前后端最核心的问答入口。
    # 这里先把 query / history / evidence / options 一次性清洗干净。
    payload = data if isinstance(data, dict) else {}
    evidence = payload.get("evidence") if isinstance(payload.get("evidence"), dict) else {}
    options = payload.get("options") if isinstance(payload.get("options"), dict) else {}

    return {
        "query": str(payload.get("query", "")).strip(),
        "session_id": str(payload.get("session_id", "")).strip() or None,
        "history": _normalize_history(payload.get("history")),
        "selected_document_ids": [
            str(item or "").strip()
            for item in payload.get("selected_document_ids", [])
            if str(item or "").strip()
        ] if isinstance(payload.get("selected_document_ids"), list) else [],
        "evidence_documents": _normalize_document_evidence(evidence.get("documents")),
        "evidence_images": _normalize_image_evidence(evidence.get("images")),
        "evidence_sensors": _normalize_sensor_evidence(evidence.get("sensors")),
        "options": options,
    }

# 从 PDF 文件对象中提取纯文本。
def _extract_pdf_text(file_storage):
    # 从 Flask file_storage 解析 PDF 全文，调用 clean_text 清洗后返回 {text, char_count, page_count}。
    pdf_data = file_storage.read()
    doc = fitz.open(stream=pdf_data, filetype="pdf")
    full_text = ""
    for page_num in range(len(doc)):
        page = doc[page_num]
        full_text += page.get_text() + "\n"
    cleaned_text = clean_text(full_text)
    return {
        "text": cleaned_text,
        "char_count": len(cleaned_text),
        "page_count": len(doc),
    }


# 从 DOCX 文件对象中提取纯文本。
def _extract_docx_text(file_storage):
    # 从 Flask file_storage 解析 DOCX 全文，按段落拼接返回 {text, char_count}。
    doc = DocxDocument(file_storage)
    text = "\n".join([para.text for para in doc.paragraphs])
    return {
        "text": text,
        "char_count": len(text),
    }


# 从 TXT 文件对象中提取纯文本。
def _extract_plain_text(file_storage):
    # 从 Flask file_storage 按 UTF-8 读取 TXT 全文，异常字符忽略，返回 {text, char_count}。
    text = file_storage.read().decode('utf-8', errors='ignore')
    return {
        "text": text,
        "char_count": len(text),
    }


# 根据文件后缀分发到对应的文本解析器。
def _extract_text_from_upload(file_storage):
    # 根据 Flask file_storage 的文件名后缀分发到 PDF/DOCX/TXT 解析器（旧版兼容接口，新版用 _extract_text_from_upload_bytes）。
    filename = str(getattr(file_storage, "filename", "") or "").lower()
    if filename.endswith('.pdf'):
        return _extract_pdf_text(file_storage)
    if filename.endswith('.docx'):
        return _extract_docx_text(file_storage)
    if filename.endswith('.txt'):
        return _extract_plain_text(file_storage)
    raise ValueError('不支持的文件格式，请上传 TXT、DOCX 或 PDF')


# 去除 Data URL 前缀，只保留纯 base64 图片内容。
def _normalize_image_base64(img_base64):
    # 去除 Data URL 前缀（如 "data:image/png;base64,"），只保留纯 base64 内容。
    # 前端传图时可能带有前缀，这里统一裁掉，只保留纯 base64 内容。
    image_data = str(img_base64 or "").strip()
    if not image_data:
        raise ValueError("Empty image data")
    if "," in image_data:
        image_data = image_data.split(",", 1)[1]
    return image_data


# 从 OpenAI 兼容响应结构中提取文本内容。
def _extract_openai_message_content(payload):
    # 从 OpenAI 兼容 API 的响应中提取文本内容。
    # 兼容两种 content 格式：{"content": "纯文本"} 或 {"content": [{"type":"text","text":"..."}]}（多模态返回）
    if not isinstance(payload, dict):
        return ""
    choices = payload.get("choices")
    if not isinstance(choices, list) or not choices:
        return ""
    message = choices[0].get("message") if isinstance(choices[0], dict) else {}
    content = message.get("content") if isinstance(message, dict) else ""
    if isinstance(content, list):
        # 多模态模型返回 content 为数组，逐个拼接
        parts = []
        for item in content:
            if isinstance(item, dict):
                text = item.get("text")
                if text:
                    parts.append(str(text))
            elif item:
                parts.append(str(item))
        return "".join(parts).strip()
    return str(content or "").strip()


# 从视觉模型返回文本中剥离出 JSON 片段。
def _extract_json_block(raw_text):
    # 视觉模型有时会把 JSON 包在 markdown 代码块里，
    # 这里负责把真正的 JSON 片段剥出来。
    text = str(raw_text or "").strip()
    if not text:
        raise ValueError("视觉模型未返回内容。")

    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\s*```$", "", text)

    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError(f"视觉模型返回内容不是 JSON：{text[:120]}")
    return text[start:end + 1]


# 把视觉模型 JSON 返回规整成统一字段。
def _normalize_vision_payload(payload):
    # 把视觉模型返回规范成系统内部统一结构：
    # keywords / summary / risk_level
    if not isinstance(payload, dict):
        raise ValueError("视觉模型返回的 JSON 不是对象。")
    #这里没写 try，所以这个异常会继续上抛到接口层（/api/image-analyze），最终被最外层的 except Exception as e 兜住，转成 500 错误返回给前端
    keywords = payload.get("keywords")
    if not isinstance(keywords, list):
        keywords = []
    normalized_keywords = []
    for item in keywords:
        text = str(item or "").strip()
        if text and text not in normalized_keywords:
            normalized_keywords.append(text)
        if len(normalized_keywords) >= 5:
            break

    summary = str(payload.get("summary") or "").strip()
    risk_level = str(payload.get("risk_level") or "未识别").strip() or "未识别"
    if risk_level not in VISION_ALLOWED_RISK_LEVELS:
        risk_level = "未识别"

    if not summary:
        summary = "、".join(normalized_keywords[:3]) or "未识别"

    return {
        "keywords": normalized_keywords,
        "summary": summary,
        "risk_level": risk_level,
    }


# 构造图片分析使用的固定提示词。
def _build_vision_prompt(image_name="image"):
    # 给视觉模型一个固定输出格式的提示词；
    # 文件名只是弱提示，真正判断仍应以画面内容为准。
    hint = str(image_name or "").strip()
    if not hint:
        return VISION_PROMPT_TEMPLATE
    return (
        f"{VISION_PROMPT_TEMPLATE}"
        f" 图片文件名为《{hint}》，文件名仅作辅助参考，优先依据画面内容判断。"
    )


# 根据图片文件名推断图片 MIME 类型。
def _guess_image_mime_type(image_name="image"):
    suffix = Path(str(image_name or "image")).suffix.lower()
    if suffix in {".jpg", ".jpeg"}:
        return "image/jpeg"
    if suffix == ".webp":
        return "image/webp"
    if suffix == ".gif":
        return "image/gif"
    return "image/png"


# 调用 OpenAI 兼容视觉接口分析图片。
def _analyze_openai_vision_base64(img_base64, image_name="image"):
    # 走 OpenAI 兼容视觉接口分析图片。
    image_data = _normalize_image_base64(img_base64)
    prompt = _build_vision_prompt(image_name)
    mime_type = _guess_image_mime_type(image_name)
    api_payload = {
        "model": config.VISION_MODEL,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:{mime_type};base64,{image_data}"},
                    },
                ],
            }
        ],
        "temperature": 0.1,
        "max_tokens": config.VISION_MAX_TOKENS,
    }

    url = config.VISION_BASE_URL.rstrip("/") + "/chat/completions"
    resp = requests.post(
        url,
        headers={
            "Authorization": f"Bearer {config.require_vision_api_key()}",
            "Content-Type": "application/json",
        },
        json=api_payload,
        timeout=(12, config.VISION_READ_TIMEOUT),
    )
    payload = resp.json() if resp.content else {}
    if not resp.ok:
        err_msg = payload.get("error") if isinstance(payload, dict) else payload
        raise RuntimeError(f"视觉模型接口错误 {resp.status_code}: {err_msg or resp.text}")

    raw_content = _extract_openai_message_content(payload)
    json_block = _extract_json_block(raw_content)
    structured = _normalize_vision_payload(json.loads(json_block))

    return {
        "provider": config.VISION_PROVIDER,
        "model": config.VISION_MODEL,
        "risk_level": structured["risk_level"],
        "keywords": structured["keywords"],
        "summary_text": (
            f"{structured['summary']}\n风险等级：{structured['risk_level']}"
            if structured["risk_level"] and structured["risk_level"] != "未识别"
            else structured["summary"]
        ),
        "result": [{"keyword": keyword} for keyword in structured["keywords"]],
        "raw_content": raw_content,
        "raw_result": payload,
    }


# 作为统一入口调用当前配置的图片分析链路。
def _analyze_image_base64(img_base64, image_name="image"):
    # 图像分析统一入口：当前只保留 OpenAI 兼容视觉链路。
    return _analyze_openai_vision_base64(img_base64, image_name)

# 延迟加载 OpenCV 依赖。
def _load_cv2():
    # 视频分析依赖 OpenCV；如果没装，直接抛出清晰错误。
    try:
        import cv2
    except Exception as exc:
        raise RuntimeError("缺少 opencv-python-headless，无法分析视频。") from exc
    return cv2


# 按配置缩放视频帧尺寸。
def _resize_video_frame(frame, cv2):
    #宽超过 960px 就等比缩到 960px 宽，高也跟着等比缩放。目的是避免把 4K 视频帧原样转 base64 发给视觉模型
    if frame is None:
        return frame
    height, width = frame.shape[:2]
    max_width = max(320, int(config.VIDEO_FRAME_MAX_WIDTH))
    if width <= max_width:
        return frame
    scale = max_width / float(width)
    new_height = max(1, int(height * scale))
    return cv2.resize(frame, (max_width, new_height), interpolation=cv2.INTER_AREA)


# 计算视频抽帧时应采样的帧位置。
def _sample_video_positions(total_frames, fps, max_frames):
    # 按固定时间间隔抽帧，尽量覆盖整段视频，而不是只看前面几帧。算要抽哪些帧
    if total_frames <= 0:
        return [0]

    max_frames = max(1, int(max_frames))
    if fps and fps > 0:
        step = max(1, int(round(float(config.VIDEO_SAMPLE_SECONDS) * fps)))
    else:
        step = max(1, total_frames // max_frames)

    positions = list(range(0, total_frames, step))
    if positions and positions[-1] != total_frames - 1:
        positions.append(total_frames - 1)
    #确保最后一帧也在名单里。按 step 跳可能刚好跳过了最后一帧，把它补进去，避免视频末尾的画面被漏掉。
    positions = sorted(set(pos for pos in positions if pos >= 0))
    if len(positions) > max_frames:
        stride = max(1, len(positions) // max_frames)
        positions = positions[::stride]
        if positions and positions[-1] != total_frames - 1:
            positions.append(total_frames - 1)
        positions = sorted(set(positions))[:max_frames]
    return positions[:max_frames]


# 把秒数格式化为视频时间标签。
def _format_video_timestamp(seconds):
    #把秒数转成可读的时间戳字符串。_format_video_timestamp(15.7)   →  "00:15"
    if seconds is None:
        return "未知时刻"
    total = max(0, int(round(float(seconds))))
    minute, second = divmod(total, 60)
    hour, minute = divmod(minute, 60)
    if hour > 0:
        return f"{hour:02d}:{minute:02d}:{second:02d}"
    return f"{minute:02d}:{second:02d}"


def _video_frame_data_url(image_base64):
    text = str(image_base64 or "").strip()
    if not text:
        return ""
    return f"data:image/jpeg;base64,{text}"


def _select_video_poster_data_url(extracted, frame_reports):
    frames = list(extracted.get("frames") or [])
    if not frames:
        return ""
    frame_by_index = {
        int(item.get("frame_index") or 0): item
        for item in frames
    }
    if frame_reports:
        preferred = frame_by_index.get(int(frame_reports[0].get("frame_index") or 0))
        if preferred:
            return _video_frame_data_url(preferred.get("image_base64"))
    return _video_frame_data_url(frames[0].get("image_base64"))


def _extract_video_poster(video_path):
    cv2 = _load_cv2()
    capture = cv2.VideoCapture(str(video_path))
    if not capture.isOpened():
        return ""

    try:
        ok, frame = capture.read()
        if not ok or frame is None:
            return ""
        frame = _resize_video_frame(frame, cv2)
        ok, buffer = cv2.imencode(
            ".jpg",
            frame,
            [int(cv2.IMWRITE_JPEG_QUALITY), int(config.VIDEO_JPEG_QUALITY)],
        )
        if not ok:
            return ""
        return _video_frame_data_url(base64.b64encode(buffer.tobytes()).decode("utf-8"))
    finally:
        capture.release()


# 从视频中抽取关键帧并转为 base64。
def _extract_video_frames(video_path):
    # 抽取视频关键帧，并把每一帧转成 base64 图片供视觉模型复用。
    cv2 = _load_cv2()
    capture = cv2.VideoCapture(str(video_path))
    if not capture.isOpened():
        raise RuntimeError("无法打开视频文件。")

    try:
        fps = float(capture.get(cv2.CAP_PROP_FPS) or 0.0)# 帧率（如 30）
        total_frames = int(capture.get(cv2.CAP_PROP_FRAME_COUNT) or 0)# 总帧数（如 5400）
        duration_s = float(total_frames / fps) if fps > 0 else 0.0# 时长秒（5400/30 = 180s）
        positions = _sample_video_positions(total_frames or 1, fps, config.VIDEO_MAX_FRAMES)

        frames = []
        for position in positions:
            capture.set(cv2.CAP_PROP_POS_FRAMES, int(position))# 跳到第 N 帧
            ok, frame = capture.read()# 读画面
            if not ok or frame is None:
                continue
            frame = _resize_video_frame(frame, cv2)# 宽>960 则等比缩小
            ok, buffer = cv2.imencode(
                ".jpg",
                frame,
                [int(cv2.IMWRITE_JPEG_QUALITY), int(config.VIDEO_JPEG_QUALITY)],
            )
            if not ok:
                continue
            image_base64 = base64.b64encode(buffer.tobytes()).decode("utf-8")# 转 base64 字符串
            timestamp_s = float(position / fps) if fps > 0 else float(len(frames))
            frames.append({
                "frame_index": int(position),
                "timestamp_s": timestamp_s,
                "image_base64": image_base64,
            })

        return {
            "fps": fps,
            "total_frames": total_frames,
            "duration_s": duration_s,
            "frames": frames,
        }
    finally:
        capture.release()


# 把逐帧分析结果汇总成视频分析响应。
def _build_video_analysis_response(video_name, extracted, frame_reports):
    # 把逐帧识别结果整理成：
    # 1. 摘要文本；
    # 2. 关键片段 evidence；
    # 3. 前端视频库需要的统计字段。
    keywords = []
    evidence = []
    detail_lines = []

    for item in frame_reports:
        item_keywords = list(item.get("keywords") or [])
        for keyword in item_keywords:
            if keyword not in keywords:
                keywords.append(keyword)
        evidence.append({
            "image_name": f"{video_name} · {item.get('timestamp_label', '未知时刻')}",
            "summary": "、".join(item_keywords[:3]) or "未识别到明显异常",
            "source_type": "video_analysis",
            "timestamp_s": round(float(item.get("timestamp_s") or 0.0), 2),
            "frame_index": int(item.get("frame_index") or 0),
        })
        detail_lines.append(
            f"- {item.get('timestamp_label', '未知时刻')}：{'、'.join(item_keywords[:3]) or '未识别到明显异常'}"
        )

    if keywords:
        issue_text = "、".join(keywords[:5])
    else:
        issue_text = "未识别到明显异常"

    summary_lines = [
        f"🎬 视频《{video_name}》分析完成。",
        f"时长：{float(extracted.get('duration_s') or 0.0):.1f}s，抽取关键帧：{len(extracted.get('frames') or [])} 张。",
        f"疑似问题：{issue_text}",
    ]
    if detail_lines:
        summary_lines.append("关键片段：")
        summary_lines.extend(detail_lines[:6])
    else:
        summary_lines.append("未发现明显异常画面，建议结合现场问题继续研判。")

    return {
        "success": True,
        "video_name": video_name,
        "duration_s": round(float(extracted.get("duration_s") or 0.0), 2),
        "fps": round(float(extracted.get("fps") or 0.0), 2),
        "total_frames": int(extracted.get("total_frames") or 0),
        "frames_extracted": len(extracted.get("frames") or []),
        "frames_matched": len(frame_reports),
        "issue_keywords": keywords[:10],
        "summary_text": "\n".join(summary_lines),
        "evidence": evidence,
    }


# 处理用户注册请求。
@app.route('/api/auth/register', methods=['POST'])
def auth_register():
    # 注册接口：接收 username + password，返回 token 和用户信息（含 library_session_id）。
    payload = request.get_json(silent=True) or {}
    try:
        result = register_user(
            username=str(payload.get("username") or "").strip(),
            password=str(payload.get("password") or ""),
        )
        user = result["user"]
        return jsonify({
            "success": True,
            "token": result["token"],
            "user": {
                "id": user["id"],
                "username": user["username"],
                "session_id": user["library_session_id"],
                "created_at": user["created_at"],
            },
        })
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400


# 处理用户登录请求。
@app.route('/api/auth/login', methods=['POST'])
def auth_login():
    # 登录接口：验密成功后将持久化文档和传感器恢复到运行时，返回 token 和用户信息。
    payload = request.get_json(silent=True) or {}
    try:
        result = login_user(
            username=str(payload.get("username") or "").strip(),
            password=str(payload.get("password") or ""),
        )
        user = result["user"]
        _hydrate_persistent_runtime_state(user)
        return jsonify({
            "success": True,
            "token": result["token"],
            "user": {
                "id": user["id"],
                "username": user["username"],
                "session_id": user["library_session_id"],
                "created_at": user["created_at"],
            },
        })
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400


# 返回当前登录用户信息。
@app.route('/api/auth/me', methods=['GET'])
def auth_me():
    # 获取当前登录用户信息（前端刷新页面后用于恢复登录态）。
    try:
        user, _token = _require_auth_user()
        _hydrate_persistent_runtime_state(user)
        return jsonify({
            "success": True,
            "user": {
                "id": user["id"],
                "username": user["username"],
                "session_id": user["library_session_id"],
                "created_at": user["created_at"],
            },
        })
    except PermissionError as exc:
        return jsonify({"error": str(exc)}), 401


# 处理退出登录请求。
@app.route('/api/auth/logout', methods=['POST'])
def auth_logout():
    # 登出：删除 token，前端清除本地存储的 auth token。
    token = _extract_auth_token()
    if token:
        delete_token(token)
    return jsonify({"success": True})


# 返回当前用户持久化保存的聊天消息。
@app.route('/api/messages/list', methods=['GET'])
def list_chat_messages():
    # 获取当前用户持久化的历史聊天消息列表。
    try:
        user, _session_id, _token = _resolve_user_session(request.args.get("session_id"))
        records = list_messages(int(user["id"]))
        return jsonify({
            "success": True,
            "messages": records,
        })
    except PermissionError as exc:
        return jsonify({"error": str(exc)}), 401


# 上传文档、建检索索引并做持久化保存。
@app.route('/api/documents/upload', methods=['POST'])
def upload_document():
    """统一文档入库接口：解析文本后建立后端向量索引，图谱改为手动生成。"""
    if 'file' not in request.files:
        return jsonify({'error': '未找到文件'}), 400

    file = request.files['file']

    try:
        user, session_id, _token = _resolve_user_session(request.form.get('session_id'))
        raw_bytes = file.read()
        parsed = _extract_text_from_upload_bytes(
            str(getattr(file, 'filename', '') or '未命名文档'),
            raw_bytes,
        )
        result = ingest_document(
            session_id=session_id,
            file_name=str(getattr(file, 'filename', '') or '未命名文档'),
            text=parsed['text'],
        )
        save_document_asset(
            user_id=int(user["id"]),
            document_id=result["document_id"],
            file_name=result["file_name"],
            raw_bytes=raw_bytes,
            text_content=parsed["text"],
            char_count=result["char_count"],
            chunk_count=result["chunk_count"],
        )
        clear_session_graph(session_id)
        build_status = mark_graph_build_pending(session_id, has_documents=True)
        return jsonify({
            'success': True,
            **result,
            "knowledge_graph": {
                "build_status": build_status,
            },
        })
    except PermissionError as e:
        return jsonify({'error': str(e)}), 401
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except RuntimeError as e:
        return jsonify({'error': str(e)}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# 上传外部三元组 JSON 并写入知识图谱。
@app.route('/api/knowledge-graph/triples/upload', methods=['POST'])
def upload_knowledge_graph_triples():
    # 上传预抽取的三元组 JSON 文件，直接导入 Neo4j（跳过 LLM 抽取步骤）。
    if 'file' not in request.files:
        return jsonify({'error': '未找到三元组文件'}), 400

    file = request.files['file']
    file_name = str(getattr(file, 'filename', '') or 'triples.json')

    try:
        _user, session_id, _token = _resolve_user_session(request.form.get('session_id'))
        raw = file.read()
        if not raw:
            return jsonify({'error': '三元组文件为空'}), 400
        payload = json.loads(raw.decode('utf-8-sig'))
        graph = import_triples_graph(session_id=session_id, payload=payload, doc_name=file_name)
        stats = graph.get("stats", {})
        return jsonify({
            "success": True,
            "session_id": str(session_id or "default"),
            "file_name": file_name,
            "node_count": int(stats.get("node_count") or len(graph.get("nodes", []))),
            "relation_count": int(stats.get("relation_count") or len(graph.get("relations", []))),
            "knowledge_graph": {
                "build_status": get_graph_build_status(session_id),
            },
        })
    except PermissionError as e:
        return jsonify({'error': str(e)}), 401
    except json.JSONDecodeError as e:
        return jsonify({'error': f'三元组 JSON 解析失败: {str(e)}'}), 400
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# 删除当前用户的文档及其检索索引。
@app.route('/api/documents/remove', methods=['POST'])
def remove_uploaded_document():
    """移除当前会话中的已入库文档。"""
    payload = request.get_json(silent=True) or {}
    document_id = str(payload.get('document_id', '')).strip()

    if not document_id:
        return jsonify({'error': 'document_id不能为空'}), 400

    try:
        user, session_id, _token = _resolve_user_session(payload.get('session_id'))
    except PermissionError as e:
        return jsonify({'error': str(e)}), 401

    removed_runtime = remove_document(session_id=session_id, document_id=document_id)
    removed_persisted = delete_document_asset(int(user["id"]), document_id)
    if not removed_runtime and not removed_persisted:
        return jsonify({'error': '未找到对应文档'}), 404
    remaining_chunks = list_session_chunks(session_id)
    clear_session_graph(session_id)
    build_status = mark_graph_build_pending(session_id, has_documents=bool(remaining_chunks))
    return jsonify({
        'success': True,
        'document_id': document_id,
        "knowledge_graph": {
            "build_status": build_status,
        },
    })


# 列出当前用户的文档库。
@app.route('/api/documents/list', methods=['GET'])
def list_uploaded_documents():
    """获取当前会话已入库文档列表。"""
    try:
        user, session_id, _token = _resolve_user_session(request.args.get('session_id'))
    except PermissionError as e:
        return jsonify({'error': str(e)}), 401

    persisted_documents = list_document_assets(int(user["id"]))
    if persisted_documents and not has_session_documents(session_id):
        hydrate_session_documents(session_id, persisted_documents)
    documents = list_session_documents(session_id)
    size_by_id = {
        str(item.get("document_id")): int(item.get("size_bytes") or 0)
        for item in persisted_documents
    }
    documents = [
        {
            **item,
            "size_bytes": size_by_id.get(str(item.get("document_id")), 0),
        }
        for item in documents
    ]
    return jsonify({
        "success": True,
        "session_id": str(session_id or "default"),
        "documents": documents,
        "document_count": len(documents),
    })


# 返回当前用户的完整或聚焦知识图谱。
@app.route('/api/knowledge-graph', methods=['GET'])
def get_knowledge_graph():
    """获取当前会话完整知识图谱。"""
    try:
        _user, session_id, _token = _resolve_user_session(request.args.get('session_id'))
    except PermissionError as e:
        return jsonify({'error': str(e)}), 401
    keyword = str(request.args.get('keyword', '')).strip()
    limit = request.args.get('limit', 160)
    try:
        limit = int(limit)
    except (TypeError, ValueError):
        limit = 160

    if keyword:
        view = query_centered_graph(session_id=session_id, keyword=keyword, limit=limit)
        stats = view.get("stats", {})
    else:
        graph = get_session_graph(session_id)
        compact_view = query_graph(graph, limit=limit)
        all_links = compact_view.get("links", []) if isinstance(compact_view, dict) else []
        safe_limit = max(50, min(int(limit or 1000), 2000))
        visible_links = all_links[:safe_limit]
        visible_uids = {str(link.get("source")) for link in visible_links} | {str(link.get("target")) for link in visible_links}
        view = {
            "nodes": [node for node in compact_view.get("nodes", []) if str(node.get("uid")) in visible_uids],
            "links": visible_links,
            "relations": visible_links,
            "stats": graph.get("stats", {}),
            "has_more": len(all_links) > len(visible_links),
            "truncated": len(all_links) > len(visible_links),
        }
        stats = graph.get("stats", {}) if isinstance(graph, dict) else {}
    return jsonify({
        "success": True,
        "session_id": str(session_id or "default"),
        "nodes": view.get("nodes", []),
        "links": view.get("links", []),
        "relations": view.get("relations", []),
        "stats": stats,
        "view_stats": {
            "node_count": len(view.get("nodes", [])),
            "relation_count": len(view.get("relations", [])),
        },
        "relation_types": stats.get("relation_types", {}),
        "query": keyword,
        "center_uid": view.get("center_uid", ""),
        "matched_uids": view.get("matched_uids", []),
        "has_more": view.get("has_more", False),
        "truncated": view.get("truncated", False),
        "from_cache": view.get("from_cache", False),
    })


# 按关键词查询当前用户的图谱子图。
@app.route('/api/knowledge-graph/query', methods=['POST'])
def query_knowledge_graph():
    """按关键词返回图谱子图。"""
    payload = request.get_json(silent=True) or {}
    try:
        _user, session_id, _token = _resolve_user_session(payload.get('session_id'))
    except PermissionError as e:
        return jsonify({'error': str(e)}), 401
    keyword = str(payload.get('keyword', '')).strip()
    try:
        limit = int(payload.get("limit") or 80)
    except (TypeError, ValueError):
        limit = 80
    try:
        depth = int(payload.get("depth") or 1)
    except (TypeError, ValueError):
        depth = 1

    view = query_centered_graph(session_id=session_id, keyword=keyword, limit=limit, depth=depth)
    return jsonify({
        "success": True,
        "session_id": str(session_id or "default"),
        "nodes": view.get("nodes", []),
        "links": view.get("links", []),
        "relations": view.get("relations", []),
        "stats": view.get("stats", {}),
        "view_stats": {
            "node_count": len(view.get("nodes", [])),
            "relation_count": len(view.get("relations", [])),
        },
        "relation_types": view.get("stats", {}).get("relation_types", {}),
        "query": keyword,
        "center_uid": view.get("center_uid", ""),
        "matched_uids": view.get("matched_uids", []),
        "has_more": view.get("has_more", False),
        "truncated": view.get("truncated", False),
        "from_cache": view.get("from_cache", False),
    })


# 触发当前用户知识图谱重建。
@app.route('/api/knowledge-graph/rebuild', methods=['POST'])
def rebuild_knowledge_graph():
    """根据当前会话已上传文档异步重建知识图谱。"""
    payload = request.get_json(silent=True) or {}
    try:
        _user, session_id, _token = _resolve_user_session(payload.get('session_id'))
    except PermissionError as e:
        return jsonify({'error': str(e)}), 401
    status = start_graph_build(session_id, list_session_chunks(session_id))
    return jsonify({
        "success": True,
        "session_id": str(session_id or "default"),
        "build_status": status,
    })


# 获取当前用户图谱构建状态。
@app.route('/api/knowledge-graph/status', methods=['GET'])
def knowledge_graph_status():
    """获取当前会话知识图谱构建状态。"""
    try:
        _user, session_id, _token = _resolve_user_session(request.args.get('session_id'))
    except PermissionError as e:
        return jsonify({'error': str(e)}), 401
    status = get_graph_build_status(session_id)
    return jsonify({
        "success": True,
        "session_id": str(session_id or "default"),
        "build_status": status,
    })


# 按节点继续展开当前用户图谱邻居。
@app.route('/api/knowledge-graph/expand', methods=['POST'])
def expand_knowledge_graph():
    """按节点展开一跳邻居。"""
    payload = request.get_json(silent=True) or {}
    try:
        _user, session_id, _token = _resolve_user_session(payload.get('session_id'))
    except PermissionError as e:
        return jsonify({'error': str(e)}), 401
    node_uid = str(payload.get('node_uid', '')).strip()
    try:
        limit = int(payload.get("limit") or 60)
    except (TypeError, ValueError):
        limit = 60
    try:
        offset = int(payload.get("offset") or 0)
    except (TypeError, ValueError):
        offset = 0
    direction = str(payload.get("direction") or "both").strip().lower()

    graph = expand_graph_neighbors(session_id=session_id, node_uid=node_uid, limit=limit, offset=offset, direction=direction)
    return jsonify({
        "success": True,
        "session_id": str(session_id or "default"),
        "node_uid": node_uid,
        "nodes": graph.get("nodes", []),
        "links": graph.get("links", []),
        "relations": graph.get("relations", []),
        "stats": graph.get("stats", {}),
        "offset": graph.get("offset", offset),
        "limit": graph.get("limit", limit),
        "has_more": graph.get("has_more", False),
        "truncated": graph.get("truncated", False),
        "from_cache": graph.get("from_cache", False),
        "returned_relation_count": graph.get("returned_relation_count", len(graph.get("relations", []))),
        "total_relation_count": graph.get("total_relation_count", len(graph.get("relations", []))),
    })


# 接收并保存当前用户的传感器数据。
@app.route('/api/sensors/push', methods=['POST'])
def push_sensor_data():
    """接收传感器接口推送的数据。"""
    payload = request.get_json(silent=True) or {}
    try:
        user, session_id, _token = _resolve_user_session(payload.get('session_id'))
    except PermissionError as e:
        return jsonify({'error': str(e)}), 401
    records = payload.get("records")
    if records is None:
        records = payload.get("sensors")
    if records is None:
        records = payload.get("data")

    result = push_session_sensors(session_id=session_id, payload=records)
    save_sensor_records(int(user["id"]), result.get("latest_records") or [])
    return jsonify({
        "success": True,
        **result,
    })


# 返回当前用户最新的传感器数据。
@app.route('/api/sensors/latest', methods=['GET'])
def latest_sensor_data():
    """获取当前会话最新传感器数据。"""
    try:
        user, session_id, _token = _resolve_user_session(request.args.get('session_id'))
    except PermissionError as e:
        return jsonify({'error': str(e)}), 401
    records = list_session_sensors(session_id)
    if not records:
        records = list_sensor_records(int(user["id"]))
        if records:
            push_session_sensors(session_id=session_id, payload=records)
    return jsonify({
        "success": True,
        "session_id": str(session_id or "default"),
        "sensor_count": len(records),
        "records": records,
    })


# 清空当前用户的传感器数据。
@app.route('/api/sensors/clear', methods=['POST'])
def clear_sensor_data():
    """清空当前会话传感器缓存。"""
    payload = request.get_json(silent=True) or {}
    try:
        user, session_id, _token = _resolve_user_session(payload.get('session_id'))
    except PermissionError as e:
        return jsonify({'error': str(e)}), 401
    cleared = clear_session_sensors(session_id)
    clear_sensor_records(int(user["id"]))
    return jsonify({
        "success": True,
        "session_id": str(session_id or "default"),
        "cleared": cleared,
    })


@app.route('/api/sensors/remove', methods=['POST'])
def remove_sensor_data():
    payload = request.get_json(silent=True) or {}
    sensor_id = str(payload.get("sensor_id") or "").strip()
    if not sensor_id:
        return jsonify({"error": "sensor_id不能为空"}), 400
    try:
        user, session_id, _token = _resolve_user_session(payload.get('session_id'))
    except PermissionError as e:
        return jsonify({'error': str(e)}), 401

    removed_runtime = remove_session_sensor(session_id, sensor_id)
    removed_persisted = delete_sensor_record(int(user["id"]), sensor_id)
    if not removed_runtime and not removed_persisted:
        return jsonify({"error": "未找到对应传感器"}), 404

    latest_records = list_session_sensors(session_id)
    if not latest_records:
        persisted_records = list_sensor_records(int(user["id"]))
        if persisted_records:
            push_session_sensors(session_id=session_id, payload=persisted_records)
            latest_records = list_session_sensors(session_id)
    return jsonify({
        "success": True,
        "sensor_id": sensor_id,
        "records": latest_records,
    })

# 上传图片并持久化保存到用户图片库。
@app.route('/api/images/upload', methods=['POST'])
def upload_persistent_image():
    # 上传图片并持久化：存储文件 + 调用视觉模型分析 + 保存摘要和 evidence。
    if 'file' not in request.files:
        return jsonify({'error': '未找到图片文件'}), 400

    file = request.files['file']
    try:
        user, _session_id, _token = _resolve_user_session(request.form.get('session_id'))
        raw_bytes = file.read()
        if not raw_bytes:
            return jsonify({'error': '图片文件为空'}), 400
        file_name = str(getattr(file, 'filename', '') or 'image.png')
        image_base64 = base64.b64encode(raw_bytes).decode("utf-8")
        try:
            analysis = _analyze_image_base64(image_base64, file_name)
        except Exception:
            analysis = {
                "summary_text": "",
                "keywords": [],
            }
        keywords = list(analysis.get("keywords") or [])[:5]
        summary_core = str(analysis.get("summary_text") or "").strip()
        evidence = []
        if keywords:
            evidence.append({
                "image_name": file_name,
                "summary": summary_core or "、".join(keywords),
                "source_type": "image_analysis",
            })
        item = save_image_asset(
            user_id=int(user["id"]),
            file_name=file_name,
            raw_bytes=raw_bytes,
            mime_type=_guess_upload_mime(file_name, "image/png"),
            summary_text=summary_core,
            evidence=evidence,
        )
        return jsonify({
            "success": True,
            **item,
        })
    except PermissionError as e:
        return jsonify({'error': str(e)}), 401
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# 列出当前用户图片库中的图片。
@app.route('/api/images/list', methods=['GET'])
def list_persistent_images():
    # 列出当前用户所有已持久化的图片，含 Data URL 和视觉分析结果。
    try:
        user, _session_id, _token = _resolve_user_session(request.args.get('session_id'))
        return jsonify({
            "success": True,
            "images": list_image_assets(int(user["id"])),
        })
    except PermissionError as e:
        return jsonify({'error': str(e)}), 401


# 从当前用户图片库中删除图片。
@app.route('/api/images/remove', methods=['POST'])
def remove_persistent_image():
    # 删除持久化图片：数据库记录和磁盘文件同时清除。
    payload = request.get_json(silent=True) or {}
    image_id = str(payload.get("image_id") or "").strip()
    if not image_id:
        return jsonify({"error": "image_id不能为空"}), 400
    try:
        user, _session_id, _token = _resolve_user_session(payload.get('session_id'))
    except PermissionError as e:
        return jsonify({'error': str(e)}), 401
    removed = delete_image_asset(int(user["id"]), image_id)
    if not removed:
        return jsonify({"error": "未找到对应图片"}), 404
    return jsonify({"success": True, "image_id": image_id})


# 上传视频并持久化保存到用户视频库。
@app.route('/api/videos/upload', methods=['POST'])
def upload_persistent_video():
    # 上传视频并持久化：存储文件 + 抽帧 + 逐帧视觉分析 + 保存 summary/evidence/keywords。
    if 'file' not in request.files:
        return jsonify({'error': '未找到视频文件'}), 400

    file = request.files['file']
    file_name = str(getattr(file, "filename", "") or "未命名视频")
    suffix = Path(file_name).suffix or ".mp4"
    temp_path = None
    try:
        user, _session_id, _token = _resolve_user_session(request.form.get('session_id'))
        raw_bytes = file.read()
        if not raw_bytes:
            return jsonify({'error': '视频文件为空'}), 400
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            temp_path = tmp.name
            tmp.write(raw_bytes)

        extracted = _extract_video_frames(Path(temp_path))
        frame_reports = []
        for frame in extracted["frames"]:
            frame_name = f"{file_name} @ {_format_video_timestamp(frame['timestamp_s'])}"
            try:
                result = _analyze_image_base64(frame["image_base64"], frame_name)
            except Exception:
                continue
            keywords = list(result.get("keywords") or [])[:5]
            if not keywords:
                continue
            frame_reports.append({
                "frame_index": frame["frame_index"],
                "timestamp_s": frame["timestamp_s"],
                "timestamp_label": _format_video_timestamp(frame["timestamp_s"]),
                "keywords": keywords,
            })

        payload = _build_video_analysis_response(file_name, extracted, frame_reports)
        item = save_video_asset(
            user_id=int(user["id"]),
            file_name=file_name,
            raw_bytes=raw_bytes,
            mime_type=_guess_upload_mime(file_name, "video/mp4"),
            poster_data_url=_select_video_poster_data_url(extracted, frame_reports),
            duration_s=float(payload.get("duration_s") or 0.0),
            frames_extracted=int(payload.get("frames_extracted") or 0),
            frames_matched=int(payload.get("frames_matched") or 0),
            summary_text=str(payload.get("summary_text") or ""),
            issue_keywords=list(payload.get("issue_keywords") or []),
            evidence=list(payload.get("evidence") or []),
        )
        return jsonify({
            "success": True,
            **item,
        })
    except PermissionError as e:
        return jsonify({'error': str(e)}), 401
    except RuntimeError as e:
        return jsonify({"error": str(e), "summary_text": "", "evidence": []}), 500
    except Exception as e:
        return jsonify({"error": str(e), "summary_text": "", "evidence": []}), 500
    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except Exception:
                pass


# 列出当前用户视频库中的视频。
@app.route('/api/videos/list', methods=['GET'])
def list_persistent_videos():
    # 列出当前用户所有已持久化的视频，含抽帧统计、关键词和 evidence。
    try:
        user, _session_id, _token = _resolve_user_session(request.args.get('session_id'))
        return jsonify({
            "success": True,
            "videos": list_video_assets(int(user["id"])),
        })
    except PermissionError as e:
        return jsonify({'error': str(e)}), 401


# 从当前用户视频库中删除视频。
@app.route('/api/videos/remove', methods=['POST'])
def remove_persistent_video():
    # 删除持久化视频：数据库记录和磁盘文件同时清除。
    payload = request.get_json(silent=True) or {}
    video_id = str(payload.get("video_id") or "").strip()
    if not video_id:
        return jsonify({"error": "video_id不能为空"}), 400
    try:
        user, _session_id, _token = _resolve_user_session(payload.get('session_id'))
    except PermissionError as e:
        return jsonify({'error': str(e)}), 401
    removed = delete_video_asset(int(user["id"]), video_id)
    if not removed:
        return jsonify({"error": "未找到对应视频"}), 404
    return jsonify({"success": True, "video_id": video_id})


# 图片分析接口（旧版兼容，无鉴权，接收 base64 返回视觉分析结果）。
# 提供单张图片分析接口。
@app.route('/api/image-analyze', methods=['POST'])
def image_analyze():
    try:
        data = request.get_json()
        
        if not data or "image_base64" not in data:
            return jsonify({"error": "No image data provided", "result": []}), 400
        
        img_base64 = data.get("image_base64", "")
        img_name = data.get("image_name", "image")
        
        if not img_base64:
            return jsonify({"error": "Empty image data", "result": []}), 400
        result = _analyze_image_base64(img_base64, img_name)
        return jsonify(result)
    except Exception as e:
        print(f"图片分析错误: {e}")
        return jsonify({"error": str(e), "result": []}), 500


# 视频分析接口（旧版兼容，无鉴权，上传视频文件返回抽帧分析结果）。
# 提供单个视频抽帧分析接口。
@app.route('/api/video-analyze', methods=['POST'])
def video_analyze():
    """处理视频上传并抽帧分析。"""
    if 'file' not in request.files:
        return jsonify({'error': '未找到视频文件'}), 400

    file = request.files['file']
    video_name = str(getattr(file, "filename", "") or "未命名视频")
    suffix = Path(video_name).suffix or ".mp4"
    temp_path = None

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            temp_path = tmp.name
        file.save(temp_path)

        extracted = _extract_video_frames(Path(temp_path))
        frame_reports = []
        for frame in extracted["frames"]:
            frame_name = f"{video_name} @ {_format_video_timestamp(frame['timestamp_s'])}"
            result = _analyze_image_base64(frame["image_base64"], frame_name)
            keywords = list(result.get("keywords") or [])[:5]
            if not keywords:
                continue
            frame_reports.append({
                "frame_index": frame["frame_index"],
                "timestamp_s": frame["timestamp_s"],
                "timestamp_label": _format_video_timestamp(frame["timestamp_s"]),
                "keywords": keywords,
            })

        payload = _build_video_analysis_response(video_name, extracted, frame_reports)
        return jsonify(payload)
    except RuntimeError as e:
        return jsonify({"error": str(e), "summary_text": "", "evidence": []}), 500
    except Exception as e:
        print(f"视频分析错误: {e}")
        return jsonify({"error": str(e), "summary_text": "", "evidence": []}), 500
    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except Exception:
                pass


# 提供普通模型代理聊天接口。
@app.route('/api/chat', methods=['POST'])
def chat_with_longcat():
    """后端代理 LongCat 请求，避免浏览器直连外网导致超时或受限。"""
    try:
        payload = request.get_json(silent=True) or {}
        system_prompt = payload.get("system", "")
        messages = payload.get("messages", [])
        model = payload.get("model") or config.LONGCAT_MODEL
        max_tokens = int(payload.get("max_tokens", 2048))

        if not isinstance(messages, list) or not messages:
            return jsonify({"error": "messages 不能为空"}), 400

        api_payload = {
            "model": model,
            "max_tokens": max_tokens,
            "messages": messages,
        }
        if system_prompt:
            api_payload["system"] = system_prompt

        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {config.require_longcat_api_key()}",
            "anthropic-version": "2023-06-01",
        }

        url = config.LONGCAT_CHAT_PROXY_URL
        # 连接超时 12 秒，读取超时使用配置值。
        resp = requests.post(
            url,
            headers=headers,
            json=api_payload,
            timeout=(12, config.LONGCAT_READ_TIMEOUT),
        )
        data = resp.json() if resp.content else {}

        if not resp.ok:
            err = data.get("error", {}) if isinstance(data, dict) else {}
            err_msg = err.get("message") or data.get("error") or resp.text
            return jsonify({"error": f"LongCat API错误 {resp.status_code}: {err_msg}"}), resp.status_code

        reply = "无响应"
        content = data.get("content", []) if isinstance(data, dict) else []
        if isinstance(content, list):
            text_block = next((b for b in content if isinstance(b, dict) and b.get("type") == "text" and b.get("text")), None)
            if text_block:
                reply = text_block.get("text", "无响应")

        return jsonify({"reply": reply, "raw": data})
    except requests.exceptions.Timeout:
        return jsonify({"error": "LongCat 请求超时（后端150秒）"}), 504
    except Exception as e:
        print(f"LongCat代理错误: {e}")
        return jsonify({"error": f"LongCat代理错误: {str(e)}"}), 500

# LangChain智能体对话接口
# 提供多智能体问答主接口。
@app.route('/api/agent-chat', methods=['POST'])
def agent_chat():
    """
    LangChain智能体问答接口
    入参兼容旧格式 {"query": "你的问题"}，也支持结构化会话与证据字段。
    出参: {"reply": "智能体回复"}
    """
    try:
        normalized = _normalize_agent_chat_request(request.get_json(silent=True) or {})
        if not normalized["query"]:
            return jsonify({"error": "query不能为空"}), 400
        user, session_id, _token = _resolve_user_session(normalized.get("session_id"))
        _hydrate_persistent_runtime_state(user)

        options = normalized["options"] or {}
        retrieval_documents = normalized["evidence_documents"]
        sensor_records = normalized["evidence_sensors"]
        use_retrieval = options.get("use_retrieval_evidence", True)
        use_sensor_evidence = options.get("use_sensor_evidence", True)

        if config.RAG_ENABLED and use_retrieval and has_session_documents(session_id):
            provisional_risk_profile = build_risk_profile(
                normalized["query"],
                normalized["history"],
                retrieval_documents,
                normalized["evidence_images"],
                sensor_records,
            )
            retrieval_documents = retrieve_relevant_chunks(
                session_id=session_id,
                query=normalized["query"],
                top_k=config.RAG_TOP_K,
                risk_types=list(provisional_risk_profile.get("risk_types", [])),
                risk_signals=list(provisional_risk_profile.get("signals_detected", [])),
                allowed_document_ids=normalized["selected_document_ids"],
            )

        if use_sensor_evidence and not sensor_records and has_session_sensors(session_id):
            sensor_records = list_session_sensors(session_id)

        result = multi_agent_ask(
            query=normalized["query"],
            session_id=session_id,
            history=normalized["history"],
            evidence_documents=retrieval_documents,
            evidence_images=normalized["evidence_images"],
            evidence_sensors=sensor_records,
            options=options,
        )
        save_message(int(user["id"]), "user", normalized["query"])
        save_message(int(user["id"]), "assistant", str(result.get("reply") or ""))
        return jsonify(result)
    except PermissionError as e:
        return jsonify({"error": str(e)}), 401
    except TimeoutError as e:
        return jsonify({"error": str(e)}), 504
    except RuntimeError as e:
        if "timed out" in str(e).lower() or "timeout" in str(e).lower() or "超时" in str(e):
            return jsonify({"error": str(e)}), 504
        return jsonify({"error": str(e)}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == '__main__':
    app.run(debug=False, port=config.SERVER_PORT)
