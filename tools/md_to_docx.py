import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def set_run_code_style(run):
    run.font.name = 'Courier New'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
    run.font.size = Pt(9)


def md_to_docx(md_path: Path, out_path: Path):
    doc = Document()
    lines = md_path.read_text(encoding='utf-8').splitlines()
    in_code = False
    code_lines = []
    table_mode = False
    table_rows = []

    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                # flush code block
                p = doc.add_paragraph()
                for cl in code_lines:
                    r = p.add_run(cl)
                    set_run_code_style(r)
                    p = doc.add_paragraph()
                in_code = False
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Tables
        if '|' in line and re.match(r'\s*\|', line):
            # collect table block
            table_mode = True
            table_rows = [line]
            j = i+1
            while j < len(lines) and '|' in lines[j]:
                table_rows.append(lines[j])
                j += 1
            # parse rows into cells
            rows = [ [c.strip() for c in re.split(r'\|', r)[1:-1]] for r in table_rows ]
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    table.cell(ri, ci).text = cell
            i = j
            continue

        # Headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            if text:
                doc.add_heading(text, level=min(level,4))
            i += 1
            continue

        # Lists
        m = re.match(r'\s*[-\*]\s+(.*)', line)
        m2 = re.match(r'\s*\d+\.\s+(.*)', line)
        if m or m2:
            txt = m.group(1) if m else m2.group(1)
            p = doc.add_paragraph(txt, style='List Bullet')
            i += 1
            continue

        # Normal paragraph
        if line.strip() == '':
            doc.add_paragraph('')
        else:
            doc.add_paragraph(line)
        i += 1

    doc.save(out_path)


if __name__ == '__main__':
    import argparse
    p = argparse.ArgumentParser()
    p.add_argument('-i','--input', required=True)
    p.add_argument('-o','--output', required=True)
    args = p.parse_args()
    md_path = Path(args.input)
    out_path = Path(args.output)
    if not md_path.exists():
        print('input not found', md_path)
        sys.exit(2)
    md_to_docx(md_path, out_path)
    print('saved', out_path)
